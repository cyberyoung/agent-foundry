#!/usr/bin/env python3
"""
Docx to Obsidian Markdown Converter

Converts .docx files into Obsidian-compatible .md notes with:
- Embedded image extraction and wiki-link embedding
- Layout table (2-col) → heading+content conversion
- Data table (3+ col) → markdown table conversion
- Rich text formatting preservation (highlight, color, bold)
- Hierarchical list detection (category headers vs sub-items)
- Chapter TOC with anchor links
- YAML frontmatter generation

Usage:
    python3 docx_to_obsidian.py <input.docx> <output-dir> [--analyze] [--no-layout-tables]

Requirements:
    pip3 install python-docx
"""

import sys
import os
import re
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is required. Install with:")
    print("  pip3 install --break-system-packages python-docx")
    sys.exit(1)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_VAL_ATTR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------


def get_run_format(run):
    """Return (is_yellow_highlight, is_red_text, is_bold) for a run."""
    is_yellow = False
    is_red = False
    is_bold = bool(run.bold)

    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        shd = rPr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            if fill and fill.upper() == "FFFF00":
                is_yellow = True

    if run.font.color and run.font.color.rgb:
        color = str(run.font.color.rgb).upper()
        if color == "FF0000":
            is_red = True

    return is_yellow, is_red, is_bold


def format_paragraph_text(para):
    """Build markdown text for a paragraph, preserving highlight/red/bold."""
    runs_info = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        is_yellow, is_red, is_bold = get_run_format(run)
        runs_info.append((text, is_yellow, is_red, is_bold))

    if not runs_info:
        return para.text  # fallback to plain text

    has_format = any(r[1] or r[2] for r in runs_info)
    if not has_format:
        return "".join(r[0] for r in runs_info)

    # Build formatted text with highlight merging
    result = []
    in_highlight = False

    for text, is_yellow, is_red, is_bold in runs_info:
        if is_yellow and not in_highlight:
            result.append("==")
            in_highlight = True
        elif not is_yellow and in_highlight:
            result.append("==")
            in_highlight = False

        if is_red:
            result.append(f'<span style="color:red">{text}</span>')
        else:
            result.append(text)

    if in_highlight:
        result.append("==")

    formatted = "".join(result)
    # Clean up trailing spaces inside highlight markers
    formatted = re.sub(r"\s+==\s*$", "==", formatted)
    return formatted


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------


def extract_images(doc, output_dir, docx_stem):
    """Extract all images from the docx and return a mapping of rId -> filename."""
    assets_dir = os.path.join(output_dir, "assets", docx_stem)
    os.makedirs(assets_dir, exist_ok=True)

    image_map = {}  # rId -> relative path for obsidian
    image_counter = 0

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_counter += 1
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            filename = f"image{image_counter}{ext}"
            filepath = os.path.join(assets_dir, filename)

            with open(filepath, "wb") as f:
                f.write(rel.target_part.blob)

            image_map[rel_id] = f"assets/{docx_stem}/{filename}"

    return image_map


def paragraph_has_image(para):
    """Check if a paragraph contains an embedded image, return rId if found."""
    for drawing in para._element.iter(qn("w:drawing")):
        for blip in drawing.iter(qn("a:blip")):
            embed = blip.get(qn("r:embed"))
            if embed:
                return embed
    return None


# ---------------------------------------------------------------------------
# Hierarchical list detection
# ---------------------------------------------------------------------------


def is_category_header(text):
    """Heuristic: short labels without stock details are category headers."""
    if not text.strip():
        return False

    text = text.strip()

    # Sub-item patterns (stock descriptions, detail lines)
    sub_patterns = [
        r"^[^\(（]{0,5}[（\(]",  # Short prefix then parenthesis
        r"^\w{2,4}股份[（\(]",  # XX股份(...)
        r"^\w{2,4}集团[（\(]",  # XX集团(...)
        r"^对标",
        r"^关注$",
        r"^受益",
        r"^深耕",
        r"^传统",
        r"^多重特性",
    ]
    for pattern in sub_patterns:
        if re.match(pattern, text):
            return False

    # Short text without parenthetical details → likely category
    if len(text) <= 10 and "(" not in text and "（" not in text:
        return True
    if text.endswith("：") or text.endswith(":"):
        return True
    if len(text) <= 15 and ("&" in text or "/" in text):
        return True
    if len(text) > 15:
        return False

    return True


def build_hierarchical_list(paragraphs):
    """Convert a list of paragraph texts into hierarchical markdown list."""
    lines = []
    for text in paragraphs:
        text = text.strip()
        if not text:
            continue
        if is_category_header(text):
            lines.append(f"- {text}")
        else:
            lines.append(f"  - {text}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Table conversion
# ---------------------------------------------------------------------------


def table_is_layout(table):
    """Determine if a table is layout (2-col label|content) or data.

    Layout tables have: exactly 2 columns, first column contains short category
    labels, second column contains longer content. The key signal is that column 1
    cells are consistently much shorter than column 2 cells.
    """
    if len(table.columns) != 2:
        return False

    col1_lengths = []
    col2_lengths = []
    for row in table.rows:
        if len(row.cells) >= 2:
            c1 = row.cells[0].text.strip()
            c2 = row.cells[1].text.strip()
            if c1 or c2:  # skip completely empty rows
                col1_lengths.append(len(c1))
                col2_lengths.append(len(c2))

    if not col1_lengths:
        return False

    avg_col1 = sum(col1_lengths) / len(col1_lengths)
    avg_col2 = sum(col2_lengths) / len(col2_lengths) if col2_lengths else 0

    # Layout table: col1 is short labels (avg < 20 chars), col2 is much longer
    # Also check: most col1 cells are short (< 30 chars)
    short_col1 = sum(1 for l in col1_lengths if l <= 30)
    return avg_col1 < 25 and short_col1 >= len(col1_lengths) * 0.7


def _format_table_cell_text(cell, image_map=None):
    """Format table cell text while preserving run-level styles."""
    parts = []
    for para in cell.paragraphs:
        img_id = paragraph_has_image(para)
        if image_map and img_id and img_id in image_map:
            parts.append(f"![[{image_map[img_id]}]]")
            continue

        text = format_paragraph_text(para).strip()
        if text:
            parts.append(text)

    if not parts:
        text = cell.text.strip().replace("\n", " ")
    else:
        text = "<br>".join(parts)

    return text.replace("|", r"\|")


def convert_data_table(table, image_map=None):
    """Convert a data table to markdown table syntax."""
    rows_data = []
    for row in table.rows:
        row_texts = [_format_table_cell_text(cell, image_map) for cell in row.cells]
        rows_data.append(row_texts)

    if not rows_data:
        return ""

    # First row as header
    header = rows_data[0]
    separator = ["---"] * len(header)
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(separator) + " |")
    for row in rows_data[1:]:
        # Pad row to match header length
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[: len(header)]) + " |")

    return "\n".join(lines)


def convert_layout_table_row(row, image_map, heading_level="####"):
    """Convert a 2-column layout table row to heading + content."""
    if len(row.cells) < 2:
        return ""

    label = row.cells[0].text.strip()
    content_cell = row.cells[1]

    if not label and not content_cell.text.strip():
        # Check for images in empty-text rows
        parts = []
        for para in content_cell.paragraphs:
            img_id = paragraph_has_image(para)
            if img_id and img_id in image_map:
                parts.append(f"\n![[{image_map[img_id]}]]")
        return "\n".join(parts) if parts else ""

    lines = []

    # Label becomes heading
    if label:
        lines.append(f"\n{heading_level} {label}\n")

    # Content: process each paragraph
    content_paragraphs = []
    for para in content_cell.paragraphs:
        img_id = paragraph_has_image(para)
        if img_id and img_id in image_map:
            # Flush accumulated text as list
            if content_paragraphs:
                lines.append(build_hierarchical_list(content_paragraphs))
                content_paragraphs = []
            lines.append(f"\n![[{image_map[img_id]}]]\n")
        else:
            text = format_paragraph_text(para)
            if text.strip():
                content_paragraphs.append(text)

    if content_paragraphs:
        # Determine if this should be a list or paragraphs
        # Use ratio of short items: if majority are short, render as list
        short_items = sum(1 for t in content_paragraphs if len(t.strip()) <= 200)
        if short_items > len(content_paragraphs) * 0.5:
            lines.append(build_hierarchical_list(content_paragraphs))
        else:
            for text in content_paragraphs:
                lines.append(text.strip())
                lines.append("")  # blank line between paragraphs

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Heading and paragraph conversion
# ---------------------------------------------------------------------------

HEADING_MAP = {
    "Heading 1": "##",
    "Heading 2": "###",
    "Heading1": "##",
    "Heading2": "###",
    "Heading 3": "####",
    "Heading3": "####",
    "Heading 4": "#####",
    "Heading4": "#####",
}


def get_heading_level(para):
    """Return markdown heading prefix if paragraph is a heading, else None."""
    style_name = para.style.name if para.style else ""
    return HEADING_MAP.get(style_name)


class NumberingResolver:
    def __init__(self, doc):
        self.num_to_abstract = {}
        self.level_map = {}
        self._build_maps(doc)

    def _build_maps(self, doc):
        numbering_part = getattr(doc.part, "numbering_part", None)
        if numbering_part is None:
            return

        numbering_element = getattr(numbering_part, "_element", None)
        if numbering_element is None:
            return

        for num in numbering_element.findall(f"{{{W_NS}}}num"):
            num_id = num.get(f"{{{W_NS}}}numId")
            abstract_num_id_elem = num.find(f"{{{W_NS}}}abstractNumId")
            if not num_id or abstract_num_id_elem is None:
                continue
            abstract_num_id = abstract_num_id_elem.get(W_VAL_ATTR)
            if abstract_num_id:
                self.num_to_abstract[str(num_id)] = str(abstract_num_id)

        for abstract_num in numbering_element.findall(f"{{{W_NS}}}abstractNum"):
            abstract_num_id = abstract_num.get(f"{{{W_NS}}}abstractNumId")
            if not abstract_num_id:
                continue

            for lvl in abstract_num.findall(f"{{{W_NS}}}lvl"):
                ilvl = lvl.get(f"{{{W_NS}}}ilvl")
                if ilvl is None:
                    continue

                num_fmt_elem = lvl.find(f"{{{W_NS}}}numFmt")
                lvl_text_elem = lvl.find(f"{{{W_NS}}}lvlText")
                num_fmt = (
                    num_fmt_elem.get(W_VAL_ATTR) if num_fmt_elem is not None else None
                )
                lvl_text = (
                    lvl_text_elem.get(W_VAL_ATTR) if lvl_text_elem is not None else ""
                )

                if num_fmt:
                    self.level_map[(str(abstract_num_id), str(ilvl))] = (
                        str(num_fmt),
                        str(lvl_text),
                    )

    def resolve(self, num_id, ilvl):
        if num_id is None or ilvl is None:
            return None

        abstract_num_id = self.num_to_abstract.get(str(num_id))
        if abstract_num_id is None:
            return None

        return self.level_map.get((str(abstract_num_id), str(ilvl)))


def get_paragraph_numbering(para_element, resolver):
    if para_element is None or resolver is None:
        return None

    ppr = para_element.find(f"{{{W_NS}}}pPr")
    if ppr is None:
        return None

    num_pr = ppr.find(f"{{{W_NS}}}numPr")
    if num_pr is None:
        return None

    num_id_elem = num_pr.find(f"{{{W_NS}}}numId")
    if num_id_elem is None:
        return None

    num_id = num_id_elem.get(W_VAL_ATTR)
    if num_id is None:
        return None

    ilvl_elem = num_pr.find(f"{{{W_NS}}}ilvl")
    ilvl_raw = ilvl_elem.get(W_VAL_ATTR) if ilvl_elem is not None else "0"

    try:
        ilvl_int = int(ilvl_raw)
    except (TypeError, ValueError):
        return None

    resolved = resolver.resolve(num_id, ilvl_int)
    if resolved is None:
        return None

    num_fmt, _lvl_text = resolved
    return num_fmt, ilvl_int, str(num_id)


# ---------------------------------------------------------------------------
# TOC generation
# ---------------------------------------------------------------------------


def slugify(text):
    """Create an anchor slug from heading text (Obsidian compatible)."""
    # Obsidian uses the heading text itself as anchor, lowercased
    slug = text.strip().lower()
    slug = re.sub(r"[^\w\u4e00-\u9fff\s-]", "", slug)  # keep Chinese chars
    slug = re.sub(r"\s+", "-", slug)
    return slug


def generate_toc(headings):
    """Generate a table of contents from collected headings.
    headings: list of (level_str, text) like ('##', 'W1'), ('###', '2026-03-01')
    """
    toc_lines = []
    for level, text in headings:
        depth = len(level) - 2  # ## = 0, ### = 1, #### = 2
        indent = "  " * depth
        anchor = text.strip()
        toc_lines.append(f"{indent}- [{anchor}](#{anchor})")
    return "\n".join(toc_lines)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------


def analyze_docx(doc_path):
    """Print structural analysis of a docx file."""
    doc = Document(doc_path)

    print(f"=== Document Analysis: {os.path.basename(doc_path)} ===\n")

    # Headings
    headings = []
    for para in doc.paragraphs:
        level = get_heading_level(para)
        if level:
            headings.append((level, para.text.strip()))
    print(f"Headings: {len(headings)}")
    for level, text in headings[:20]:
        print(f"  {level} {text}")
    if len(headings) > 20:
        print(f"  ... and {len(headings) - 20} more")

    # Tables
    print(f"\nTables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables[:10]):
        rows = len(table.rows)
        cols = len(table.columns)
        layout = "layout" if table_is_layout(table) else "data"
        first_cell = table.rows[0].cells[0].text.strip()[:30] if table.rows else ""
        print(f"  Table {i}: {rows} rows x {cols} cols ({layout}) — '{first_cell}'")

    # Images
    image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    print(f"\nImages: {image_count}")

    # Formatting
    yellow = 0
    red = 0
    bold = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        is_y, is_r, is_b = get_run_format(run)
                        if is_y:
                            yellow += 1
                        if is_r:
                            red += 1
                        if is_b:
                            bold += 1
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            is_y, is_r, is_b = get_run_format(run)
            if is_y:
                yellow += 1
            if is_r:
                red += 1
            if is_b:
                bold += 1

    print(f"\nFormatting runs:")
    print(f"  Yellow highlight: {yellow}")
    print(f"  Red text: {red}")
    print(f"  Bold: {bold}")


def convert_docx(doc_path, output_dir, use_layout_tables=True):
    """Main conversion: docx → Obsidian markdown + extracted images."""
    doc = Document(doc_path)
    docx_stem = Path(doc_path).stem
    # Strip common export suffixes (e.g. "xxx.tencentdoc.docx" → "xxx")
    for suffix in [".tencentdoc"]:
        if docx_stem.endswith(suffix):
            docx_stem = docx_stem[: -len(suffix)]

    # Extract images
    image_map = extract_images(doc, output_dir, docx_stem)
    print(f"Extracted {len(image_map)} images to assets/{docx_stem}/")

    # Build document body
    # Strategy: iterate through the docx body in order (paragraphs + tables interleaved)
    body = doc.element.body
    md_lines = []
    headings_for_toc = []
    table_index = 0
    tables_by_element = {table._element: table for table in doc.tables}
    numbering_resolver = NumberingResolver(doc)
    counters = {}
    chinese_counter_map = {
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九",
        10: "十",
    }

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Paragraph (heading or body text)
            from docx.text.paragraph import Paragraph

            para = Paragraph(child, doc)
            numbering_info = get_paragraph_numbering(para._element, numbering_resolver)

            if numbering_info:
                num_fmt, ilvl, num_id = numbering_info
                text = format_paragraph_text(para).strip()
                if text:
                    counters[num_id] = counters.get(num_id, 0) + 1

                    if num_fmt == "chineseCountingThousand" and ilvl == 0:
                        chinese_num = chinese_counter_map.get(
                            counters[num_id], str(counters[num_id])
                        )
                        md_lines.append(f"\n#### {chinese_num}、{text}\n")
                    elif num_fmt == "bullet":
                        md_lines.append(f"{'  ' * ilvl}- {text}")
                    elif num_fmt == "decimal":
                        md_lines.append(f"{'  ' * ilvl}{counters[num_id]}. {text}")
                    else:
                        md_lines.append(f"{'  ' * ilvl}- {text}")
                continue

            heading_level = get_heading_level(para)
            if heading_level:
                text = para.text.strip()
                if text:
                    headings_for_toc.append((heading_level, text))
                    md_lines.append(f"\n{heading_level} {text}\n")
            else:
                # Check for image
                img_id = paragraph_has_image(para)
                if img_id and img_id in image_map:
                    md_lines.append(f"\n![[{image_map[img_id]}]]\n")
                else:
                    text = format_paragraph_text(para)
                    if text.strip():
                        md_lines.append(text.strip())
                        md_lines.append("")

        elif tag == "tbl":
            # Table
            if child in tables_by_element:
                table = tables_by_element[child]

                if use_layout_tables and table_is_layout(table):
                    # Convert layout table: each row → heading + content
                    for row in table.rows:
                        row_md = convert_layout_table_row(row, image_map)
                        if row_md.strip():
                            md_lines.append(row_md)
                else:
                    # Convert as data table
                    md_lines.append("")
                    md_lines.append(convert_data_table(table, image_map=image_map))
                    md_lines.append("")

    # Build TOC
    toc = generate_toc(headings_for_toc) if headings_for_toc else ""

    # Build frontmatter
    title = docx_stem
    # Try to get title from docx core properties
    try:
        if doc.core_properties.title:
            title = doc.core_properties.title
    except Exception:
        pass

    date_str = datetime.now().strftime("%Y-%m-%d")
    try:
        if doc.core_properties.created:
            date_str = doc.core_properties.created.strftime("%Y-%m-%d")
    except Exception:
        pass

    frontmatter = f"""---
title: {title}
date: {date_str}
tags: []
category: {os.path.basename(output_dir)}
cssclasses:
  - table-nowrap
---"""

    # Assemble final document
    parts = [frontmatter, "", f"# {title}", ""]
    if toc:
        parts.extend(["## 目录", "", toc, "", "---", ""])
    parts.extend(md_lines)

    content = "\n".join(parts)

    content = re.sub(r"\n{3,}", "\n\n", content)

    label_re = re.compile(r"^.{1,20}[：:]\s*$")
    lines = content.split("\n")
    compact = []
    i = 0
    while i < len(lines):
        if (
            i + 2 < len(lines)
            and label_re.match(lines[i])
            and lines[i + 1] == ""
            and label_re.match(lines[i + 2])
        ):
            compact.append(lines[i])
            i += 2
        else:
            compact.append(lines[i])
            i += 1
    content = "\n".join(compact)

    # Write output
    output_path = os.path.join(output_dir, f"{docx_stem}.md")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"Written: {output_path}")
    print(f"Total headings in TOC: {len(headings_for_toc)}")

    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Convert .docx to Obsidian-compatible Markdown"
    )
    parser.add_argument("input", help="Path to the .docx file")
    parser.add_argument("output_dir", help="Output directory for the .md file")
    parser.add_argument(
        "--analyze",
        action="store_true",
        help="Only analyze the docx structure, do not convert",
    )
    parser.add_argument(
        "--no-layout-tables",
        action="store_true",
        help="Treat all tables as data tables (markdown table format)",
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: File not found: {args.input}")
        sys.exit(1)

    if args.analyze:
        analyze_docx(args.input)
    else:
        os.makedirs(args.output_dir, exist_ok=True)
        convert_docx(
            args.input, args.output_dir, use_layout_tables=not args.no_layout_tables
        )


if __name__ == "__main__":
    main()
